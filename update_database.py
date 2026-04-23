"""
╔══════════════════════════════════════════════════════════════════════════╗
║          СКРИПТ ОБНОВЛЕНИЯ БАЗЫ ДАННЫХ — КМГ НДО                       ║
║                                                                          ║
║  Что делает:                                                             ║
║  1. Читает Word-файлы мониторинга из SharePoint                         ║
║  2. Обновляет лист "Monitoring" в database.xlsx                         ║
║  3. Конвертирует database.xlsx → database.json (для PWA)                ║
║  4. Пушит оба файла на GitHub                                            ║
║                                                                          ║
║  Запуск: двойной клик или  py update_database.py                        ║
╚══════════════════════════════════════════════════════════════════════════╝
"""

import os
import re
import sys
import json
import shutil
import builtins
import subprocess
from datetime import datetime, date

import pandas as pd
import openpyxl
from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn

if 'exceptions' not in sys.modules:
    sys.modules['exceptions'] = builtins

# Фикс кодировки для Windows терминала
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# ─── Пути ───────────────────────────────────────────────────────────────────

REPO_DIR       = os.path.dirname(os.path.abspath(__file__))
DATABASE_XLSX  = os.path.join(REPO_DIR, 'database.xlsx')
DATABASE_JSON  = os.path.join(REPO_DIR, 'database.json')

# Разбивка на 2 файла для ускорения загрузки на merger.kz:
#  - db_main.json: ВСЁ кроме 'msp' (~10 МБ, грузится при старте)
#  - db_msp.json:  только 'msp'   (~16 МБ, ленивая загрузка с прогресс-баром)
DB_MAIN_JSON   = os.path.join(REPO_DIR, 'db_main.json')
DB_MSP_JSON    = os.path.join(REPO_DIR, 'db_msp.json')

# Секции, которые идут в db_msp.json. Всё остальное → db_main.json.
MSP_ONLY_SECTIONS = {'msp'}

# Источник актуального Excel (из SharePoint/OneDrive)
SOURCE_XLSX    = r'C:\Users\Daryn\SharePoint\Департамент тарифной политики - Инф\Daryn\Дарын\PBI\database.xlsx'

# Папка с Word-файлами мониторинга
MONITORING_DIR = r'C:\Users\Daryn\SharePoint\Департамент тарифной политики - Инф\Daryn\Дарын\PBI\Мониторинг'

# GitHub токен — читается из окружения или из .env файла (gitignored)
def _load_github_token() -> str:
    """Читаем токен из env var GITHUB_TOKEN или из файла .env рядом со скриптом."""
    tok = os.environ.get('GITHUB_TOKEN', '').strip()
    if tok:
        return tok
    env_file = os.path.join(REPO_DIR, '.env')
    if os.path.exists(env_file):
        with open(env_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line.startswith('GITHUB_TOKEN='):
                    return line.split('=', 1)[1].strip().strip('"').strip("'")
    print('  ✗ GITHUB_TOKEN не найден. Добавь его в .env файл:')
    print(f'     echo GITHUB_TOKEN=ghp_xxx > "{env_file}"')
    return ''

GITHUB_TOKEN   = _load_github_token()
GITHUB_REPO    = 'https://github.com/DarynZhamil/VideoforPBI.git'

# ─── Шаг 1: Парсинг Word-файлов мониторинга ────────────────────────────────

def get_risk_level(cell):
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shadings = tcPr.xpath('.//w:shd')
        if shadings:
            fill = shadings[0].get(qn('w:fill')) or shadings[0].get('fill')
            if fill and len(fill) == 6:
                r, g, b = int(fill[0:2], 16), int(fill[2:4], 16), int(fill[4:6], 16)
                if (r, g, b) == (255, 0, 0):     return ("Высокая", 3)
                if (r, g, b) == (255, 255, 0):   return ("Средняя", 2)
                if (r, g, b) == (146, 208, 80):  return ("Незначительная", 1)
    except Exception:
        pass
    return ("Разрешенные", 0)


def extract_full_text(cell):
    if not cell or not cell.text:
        return ""
    # Preserve paragraph boundaries with newlines
    return '\n'.join(
        ''.join(run.text for run in p.runs).strip()
        for p in cell.paragraphs
    ).strip()


def extract_workers_count(text):
    if pd.isna(text) or text == '':
        return None
    match = re.search(r'([\d\s]+)', str(text).replace('\u202f', ' ').replace('\xa0', ' '))
    if not match:
        return None
    num = match.group(1).replace(' ', '')
    return int(num) if num.isdigit() else None


def extract_profsoyuz(text):
    if pd.isna(text) or text == '':
        return "не определен"
    block = re.search(r'(?:Профсоюз[ы]?:)(.*?)(?=\n\w+:|$)', text, re.DOTALL | re.IGNORECASE)
    if not block:
        return "не определен"
    result = re.sub(r'\s+', ' ', block.group(1)).strip()
    result = re.sub(r'\d+\.\s*', '\n', result).strip()
    return result if result else "не определен"


def extract_data_from_docx(docx_path):
    doc = Document(docx_path)
    basename = os.path.basename(docx_path)
    date_match = re.search(r'\d{2}\.\d{2}\.\d{4}', basename)
    file_date = datetime.strptime(date_match.group(0), '%d.%m.%Y').date() if date_match else None

    data = []
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 4:
                continue
            texts = [extract_full_text(c) for c in row.cells]
            risk_level, risk_code = get_risk_level(row.cells[0])
            org_name = texts[1].split('\n')[0].strip()

            details = {}
            for line in texts[1].split('\n'):
                if 'Вид деят-ти:' in line:
                    details['service_type'] = line.split(':', 1)[-1].strip()
                elif 'Кол-во работников:' in line:
                    details['employees'] = extract_workers_count(line.split(':', 1)[-1].strip())
                elif 'Заказчик:' in line:
                    details['dzo'] = line.split(':', 1)[-1].strip()

            data.append({
                'date':           file_date,
                'risk':           risk_level,
                'code':           risk_code,
                'contractor':     normalize_org_name(org_name),
                'service_type':   details.get('service_type'),
                'employees':      details.get('employees', 0) or 0,
                'average_salary': None,
                'dzo':            details.get('dzo'),
                'union':          extract_profsoyuz(texts[1]),
                'requirements':   texts[2],
                'measures':       texts[3],
                'file_name':      basename.split(' ')[0],
            })
    return data


REPLACEMENTS = [
    ("АКТАУСКИЙ ФИЛИАЛ ТОО \"KMG-SECURITY\"",         "ТОО \"KMG-SECURITY\""),
    ("ЖАНАОЗЕНСКИЙ ФИЛИАЛ ТОО \"KMG-SECURITY\"",      "ТОО \"KMG-SECURITY\""),
    ("ЖФ ТОО \"KMG-SECURITY\"",                        "ТОО \"KMG-SECURITY\""),
    ("МАНГИСТАУСКИЙ ФИЛИАЛ ТОО \"KMG-SECURITY\"",      "ТОО \"KMG-SECURITY\""),
    ("МФ ТОО \"KMG-SECURITY\"",                        "ТОО \"KMG-SECURITY\""),
    ("МАНИСТАУСКИЙ ФИЛИАЛ ТОО \"KMG-SECURITY\"",       "ТОО \"KMG-SECURITY\""),
    ("ЖФ ТОО \"СЕМСЕР – ӨРТ СӨНДІРУШІ\"",             "ТОО \"СЕМСЕР - ӨРТ СӨНДІРУШІ\""),
    ("ЖФ ТОО \"СЕМСЕР ӨРТ-СӨНДІРУШІ\"",               "ТОО \"СЕМСЕР - ӨРТ СӨНДІРУШІ\""),
    ("ТОО \"СЕМСЕР-ӨРТ СӨНДІРУШІ\"",                  "ТОО \"СЕМСЕР - ӨРТ СӨНДІРУШІ\""),
    ("ТОО \"СЕМСЕР – ӨРТ СӨНДІРУШІ\"",                "ТОО \"СЕМСЕР - ӨРТ СӨНДІРУШІ\""),
    ("ИП \"ДЖАНИБЕКОВ Н.А.\"",                         "ИП \" ДЖАНИБЕКОВ НУРБЕРГЕН АЛДАБЕРГЕНОВИЧ\""),
    ("ИП \"ҚОЙШЫБАЕВ Ш.С.\"",                          "ИП \"КОЙШЫБАЕВ Ш.С.\""),
    ("ИП \"РАТАЙ М.К.\"",                              "ИП \"РАТАЙ МАРАТ КАРАСАЙУЛЫ\""),
    ("ИП \"ТУМАНШАЕВА К.К.\"",                         "ИП ТУМАНШАЕВА"),
    ("МФ ТОО \"М-ТЕХСЕРВИС\"",                         "ТОО \"М-ТЕХСЕРВИС\" (MUNAI WELL SERVICE)"),
    ("ТОО \"DEMEY I KO\"",                             "ТОО \"DEMEY I KO"),
    ("ТОО \"NUR CASPIAN SERVICES\"",                   "ТОО \"NUR-CASPIAN SERVICES\""),
    ("ТОО \"TESTCOM OZEN\"",                           "ТОО TESTCOM OZEN"),
    ("ТОО \"АМИРОН ОЙЛ СЕРВИС\"",                     "ТОО \"АМИРОНОЙЛСЕРВИС\""),
    ("ТОО \"БАТЫС ГЕОФИЗ СЕРВИС\"",                    "ТОО \"БАТЫСГЕОФИЗСЕРВИС\""),
    ("ТОО \"БК \"ВЕЛИКАЯ СТЕНА\"",                     "ТОО БК \"ВЕЛИКАЯ СТЕНА\" (АКТАУ)"),
    ("ТОО \"БОЗАШЫ ТРАНС КУРЫЛЫС\"",                   "ТОО \"БОЗАШЫТРАНСҚҰРЫЛЫС\""),
    ("ТОО \"БОЗАШЫ ТРАНС ҚҰРЫЛЫС\"",                   "ТОО \"БОЗАШЫТРАНСҚҰРЫЛЫС\""),
    ("ТОО \"БОЛАШАК ОЗЕН\"",                           "ТОО \"БОЛАШАК-УЗЕНЬ\""),
    ("ТОО \"МЕДИКЕР – ПРОМЫШЛЕННАЯ",                   "ТОО \"МЕДИКЕР - ПРОМЫШЛЕННАЯ МЕДИЦИНА\""),
    ("ТОО \"МОБИЛСЕРВИСГРУПП\"",                       "ТОО \"МОБИЛ СЕРВИС ГРУПП\""),
    ("ТОО \"СТРОЙ МАСТЕР\"",                           "ТОО \"СТРОЙМАСТЕР\""),
    ("ТОО \"ТАМШАЛЫ-СЕРВИС\" (МАНГИСТАУСКАЯ ОБЛАСТЬ)", "ТОО \"ТАМШАЛЫ СЕРВИС\""),
    ("ТОО \"ТЕХНО ТРЕЙДИНГ ЛТД\"",                     "ТОО \"TECHNO TRADING LTD\""),
    ("ТОО \"ФИРМА \"ЧА-КУР\"",                         "ТОО \"ФИРМА ЧА-КУР\""),
    ("ТОО \"ЮГО-ВОСТОЧНАЯ СЕРВИСНАЯ ГРУППА\"",         "ТОО \"ЮВСГ\""),
    ("МФ ТОО \"РЦШПВАСС\" \"АКБЕРЕН\"",               "МФ \"АК БЕРЕН\" ТОО \"РЦШ ПВАСС\""),
    ("ТОО \"МӨЛДІРМҰНАЙСЕРВИС\"",                      "ТОО \"МӨЛДІР МҰНАЙ СЕРВИС\""),
    ("ИП \"БУРКИТАЛИЕВ Б.Н.\"",                        "ИП \"БУРКИТАЛИЕВ А.\""),
    ("ТОО \"CITIC - ВОДНАЯ ЭКОЛОГИЯ\"",                "ТОО \"СИТИК-ВОДНАЯ ЭКОЛОГИЯ\""),
]


def normalize_org_name(name: str) -> str:
    if not isinstance(name, str):
        return name
    name = name.strip().replace('«', '"').replace('»', '"').upper()
    for old, new in REPLACEMENTS:
        name = name.replace(old.upper(), new.upper())
    return name.strip()


def build_monitoring_df():
    if not os.path.isdir(MONITORING_DIR):
        print(f'  ⚠ Папка мониторинга не найдена: {MONITORING_DIR}')
        return pd.DataFrame()

    all_data = []
    files = [f for f in os.listdir(MONITORING_DIR) if f.lower().endswith('.docx')]
    print(f'  Найдено Word-файлов: {len(files)}')

    for filename in files:
        try:
            rows = extract_data_from_docx(os.path.join(MONITORING_DIR, filename))
            all_data.extend(rows)
            print(f'    ✓ {filename} ({len(rows)} записей)')
        except Exception as e:
            print(f'    ✗ {filename}: {e}')

    if not all_data:
        return pd.DataFrame()

    df = pd.DataFrame(all_data)
    df = df[df['contractor'].str.strip() != '']
    # Фильтруем строки-заголовки из Word-таблиц
    header_mask = df['contractor'].str.contains(
        'НАИМЕНОВАНИЕ ОРГАНИЗАЦИИ|НАИМЕНОВАНИЕ ПО|№ П/П|N П/П',
        case=False, na=False, regex=True
    )
    df = df[~header_mask]
    df = df.reset_index(drop=True)
    return df


# ─── Шаг 2: Обновление листа Monitoring в database.xlsx ────────────────────

def update_monitoring_sheet(df: pd.DataFrame):
    print('\n[2] Обновление листа Monitoring в database.xlsx...')

    # Копируем актуальный Excel из SharePoint если он новее
    if os.path.exists(SOURCE_XLSX):
        src_mtime = os.path.getmtime(SOURCE_XLSX)
        dst_mtime = os.path.getmtime(DATABASE_XLSX) if os.path.exists(DATABASE_XLSX) else 0
        if src_mtime > dst_mtime:
            print(f'  Копирую обновлённый database.xlsx из SharePoint...')
            shutil.copy2(SOURCE_XLSX, DATABASE_XLSX)
            print(f'  ✓ Скопирован')
        else:
            print(f'  database.xlsx актуален (SharePoint не новее)')
    else:
        print(f'  ⚠ Источник не найден: {SOURCE_XLSX}')
        if not os.path.exists(DATABASE_XLSX):
            print('  ✗ database.xlsx отсутствует — прерываю')
            return False

    if df.empty:
        print('  ⚠ Нет данных мониторинга — лист не обновляем')
        return True

    wb = load_workbook(DATABASE_XLSX)

    # Удаляем старый лист и создаём новый
    if 'Monitoring' in wb.sheetnames:
        del wb['Monitoring']
    ws = wb.create_sheet('Monitoring')

    # Заголовки
    cols = ['date', 'risk', 'code', 'contractor', 'service_type',
            'employees', 'average_salary', 'dzo', 'union', 'requirements', 'measures', 'file_name']
    ws.append(cols)

    # Данные
    for _, row in df.iterrows():
        ws.append([row.get(c) for c in cols])

    wb.save(DATABASE_XLSX)
    print(f'  ✓ Лист Monitoring обновлён ({len(df)} записей)')
    return True


# ─── Шаг 3: Конвертация database.xlsx → database.json ──────────────────────

SKIP_SHEETS = {'Медиана_ЗП', 'Медиана_ЗП_agg', 'Сумма_договоров_по_годам'}
# Mapp_prof и ЗП_должность НУЖНЫ для fallback страницы Алгоритм на merger.kz

# Для больших листов оставляем только нужные колонки (как в Flutter-приложении)
KEEP_COLS = {
    'msp':            {'dzo', 'contractor', 'net_salary', 'total_per_worker',
                       'Наименование должностей', 'Месяц', 'Месяц№'},
    'Данные договор': {'contractor', 'position_title', 'total_per_worker', 'net_salary',
                       'dzo', 'year', 'status', 'salary', 'total_employees',
                       'fzp', 'bonuses', 'group', 'compound'},
}

ROW_FILTERS = {}  # Больше не нужны — Медиана_ЗП обрабатывается отдельно


# ─── Агрегация Медиана_ЗП (228K строк → ~170 строк) ────────────────────────

_MZ_ATTRS    = ['ЗП', 'ЗП на руки', 'МСП', 'Оклад по тарифу']
_MZ_YEARS    = [2024, 2025, 2026]
_MZ_STATUSES = {'действует', 'в работе'}


def aggregate_median_zp(wb):
    """Читает лист Медиана_ЗП, фильтрует год=2026 + активные статусы,
    агрегирует по dzo × contractor × attribute → mean(value).

    Итог: ~несколько сотен строк вместо 227k.
    Структура: [{dzo, contractor, attribute, value}, ...]
    """
    if 'Медиана_ЗП' not in wb.sheetnames:
        print('  ! Медиана_ЗП: лист не найден, пропускаем')
        return []

    ws = wb['Медиана_ЗП']
    raw_rows = list(ws.iter_rows(values_only=True))
    if len(raw_rows) < 2:
        return []

    headers = [str(h) if h is not None else f'col_{i}'
               for i, h in enumerate(raw_rows[0])]

    total = len(raw_rows) - 1

    # Шаг 1: собираем строки после фильтрации (без агрегации)
    # Каждая строка = 1 сотрудник; для МСП value = фонд на всю позицию
    # (строки дублируются, поэтому row_count_per_position = employee_count)
    filtered: list[dict] = []
    for row in raw_rows[1:]:
        r = dict(zip(headers, row))

        # Фильтр по году (2026)
        dt = r.get('date')
        if dt is None:
            continue
        yr = dt.year if isinstance(dt, (datetime, date)) else None
        if yr is None:
            try:
                yr = int(str(dt)[:4])
            except (ValueError, TypeError):
                continue
        if yr != 2026:
            continue

        # Фильтр по статусу
        st = str(r.get('status') or '').strip().lower()
        if st not in _MZ_STATUSES:
            continue

        # Фильтр по атрибуту
        attr = str(r.get('attribute') or '').strip()
        if attr not in _MZ_ATTRS:
            continue

        try:
            val = float(r.get('value') or 0)
        except (ValueError, TypeError):
            continue
        if val <= 0:
            continue

        dzo        = str(r.get('dzo')            or '').strip()
        contractor = str(r.get('contractor')     or '').strip()
        position   = str(r.get('position_title') or '').strip()
        if not dzo or not contractor:
            continue

        # total_employees (n) — кол-во сотрудников на позиции; используется ТОЛЬКО для МСП
        try:
            n_emp = int(float(r.get('total_employees') or 0))
        except (ValueError, TypeError):
            n_emp = 0

        filtered.append({
            'dzo': dzo, 'contractor': contractor,
            'attr': attr, 'position': position,
            'value': val, 'n_emp': n_emp,
        })

    # Шаг 2: для МСП — нормируем value / total_employees / 12
    #   value = годовой фонд на ВСЮ позицию (на всех n сотрудников)
    #   результат = месячная зарплата на 1 сотрудника
    #
    # Дедупликация по (dzo, contractor, position_title): одна позиция — одно значение.
    # Это обеспечивает РАВНЫЙ вес каждой должности при усреднении по подрядчику,
    # независимо от количества сотрудников (n).
    #
    # Для ЗП на руки и Оклад по тарифу — value уже на 1 сотрудника, нормировка не нужна.

    # Шаг 3: строим (dzo, contractor, attribute) → [per_position_monthly_value, ...]
    #   Ключ = (dzo, contractor, attr, position) чтобы дедуплицировать позиции.
    pos_vals: dict[tuple, float] = {}   # (dzo, contractor, attr, position) → monthly value

    for r in filtered:
        dzo, contractor, attr, position, val, n_emp = (
            r['dzo'], r['contractor'], r['attr'],
            r['position'], r['value'], r['n_emp'],
        )
        if attr == 'МСП':
            if n_emp <= 0:
                continue  # нет данных о сотрудниках — пропускаем строку
            monthly = val / n_emp / 12  # годовой фонд → месячное на 1 сотрудника
        else:
            monthly = val  # ЗП на руки / Оклад — уже на 1 сотрудника

        pos_key = (dzo, contractor, attr, position)
        # Если позиция встречается несколько раз (дублированные строки) — берём среднее
        if pos_key in pos_vals:
            pos_vals[pos_key] = (pos_vals[pos_key] + monthly) / 2
        else:
            pos_vals[pos_key] = monthly

    # Шаг 4: группируем по (dzo, contractor, attr) → список значений по должностям
    groups: dict[tuple, list] = {}
    for (dzo, contractor, attr, _position), monthly in pos_vals.items():
        key = (dzo, contractor, attr)
        groups.setdefault(key, []).append(monthly)

    # Шаг 5: среднее по должностям → одно значение на подрядчика
    contractor_vals: dict[tuple, float] = {
        key: sum(vals) / len(vals)
        for key, vals in groups.items()
    }

    # Шаг 6: группируем по (dzo, attr) → среднее по подрядчикам (равный вес)
    dzo_groups: dict[tuple, list] = {}
    for (dzo, contractor, attr), avg in contractor_vals.items():
        dzo_groups.setdefault((dzo, attr), []).append((contractor, avg))

    # Результат: оставляем уровень dzo × contractor (для фильтрации по ДЗО в FastAPI)
    result = [
        {
            'dzo':        dzo,
            'contractor': contractor,
            'attribute':  attr,
            'value':      round(avg),
        }
        for (dzo, contractor, attr), avg in sorted(contractor_vals.items())
    ]

    print(f'  ✓ Медиана_ЗП_agg: {len(result)} строк (агрегировано из {total})')
    return result


def write_mediana_agg_sheet(agg_rows: list):
    """Записывает компактную таблицу Медиана_ЗП_agg в database.xlsx."""
    if not agg_rows:
        return
    wb = load_workbook(DATABASE_XLSX)
    if 'Медиана_ЗП_agg' in wb.sheetnames:
        del wb['Медиана_ЗП_agg']
    ws = wb.create_sheet('Медиана_ЗП_agg')
    ws.append(['dzo', 'contractor', 'attribute', 'value'])
    for r in agg_rows:
        ws.append([r['dzo'], r['contractor'], r['attribute'], r['value']])
    wb.save(DATABASE_XLSX)
    print(f'  ✓ Лист Медиана_ЗП_agg записан в database.xlsx ({len(agg_rows)} строк)')


def compute_median_zp(xlsx_path: str) -> list:
    """Вычисляет median_zp_computed из листа Медиана_ЗП.

    Берёт ТОЛЬКО attribute='ЗП' (не МСП/Оклад — там другие единицы!).
    Статус-правило: 2026 → {Действует, В работе}; 2024/2025 → добавляется 'Завершен'.

    Результат: [{dzo, year, attribute='ЗП', median, count}, ...]
    Формат совместим с ExcelDataService.getDzoKpi() на merger.kz.
    """
    try:
        mz = pd.read_excel(xlsx_path, sheet_name='Медиана_ЗП')
    except Exception as e:
        print(f'  ! median_zp_computed: не удалось прочитать Медиана_ЗП: {e}')
        return []

    mz['_year'] = pd.to_datetime(mz['date'], errors='coerce').dt.year
    mz['_val']  = pd.to_numeric(mz['value'], errors='coerce')

    _active = {'Действует', 'В работе'}
    _hist   = {'Действует', 'В работе', 'Завершен'}

    base = mz[
        (mz['_val'] > 0) &
        (mz['attribute'].astype(str).str.strip() == 'ЗП')
    ].copy()

    agg_rows = []
    for year in sorted(base['_year'].dropna().unique()):
        allowed = _active if int(year) >= 2026 else _hist
        yr_data = base[
            (base['_year'] == year) &
            (base['status'].astype(str).str.strip().isin(allowed))
        ]
        if yr_data.empty:
            continue

        # По каждому ДЗО
        for dzo_name, dzo_grp in yr_data.groupby('dzo'):
            vals = dzo_grp['_val'].dropna()
            vals = vals[vals > 0]
            if not vals.empty:
                agg_rows.append({
                    'dzo':       str(dzo_name).strip(),
                    'year':      int(year),
                    'attribute': 'ЗП',
                    'median':    round(float(vals.median()), 2),
                    'count':     len(vals),
                })

        # По всем ДЗО
        vals = yr_data['_val'].dropna()
        vals = vals[vals > 0]
        if not vals.empty:
            agg_rows.append({
                'dzo':       'Все',
                'year':      int(year),
                'attribute': 'ЗП',
                'median':    round(float(vals.median()), 2),
                'count':     len(vals),
            })

    print(f'  ✓ median_zp_computed: {len(agg_rows)} строк')
    return agg_rows


def xlsx_to_json():
    print('\n[3] Конвертация database.xlsx → database.json...')
    wb = load_workbook(DATABASE_XLSX, read_only=True, data_only=True)

    result = {}
    for sheet_name in wb.sheetnames:
        if sheet_name in SKIP_SHEETS:
            continue
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 2:
            continue

        headers = [str(h) if h is not None else f'col_{i}' for i, h in enumerate(rows[0])]
        keep = KEEP_COLS.get(sheet_name)  # None = оставить все колонки
        sheet_data = []
        for row in rows[1:]:
            record = {}
            has_value = False
            for h, val in zip(headers, row):
                if keep and h not in keep:
                    continue
                if isinstance(val, (datetime, date)):
                    val = val.strftime('%Y-%m-%d')
                if val is not None and str(val).strip() != '':
                    has_value = True
                record[h] = val
            if has_value:
                sheet_data.append(record)

        if sheet_data:
            result[sheet_name] = sheet_data
            print(f'  ✓ {sheet_name}: {len(sheet_data)} строк')

    # Агрегируем Медиана_ЗП отдельно (228K строк → компактная таблица)
    wb2 = load_workbook(DATABASE_XLSX, read_only=False, data_only=True)
    mz_agg = aggregate_median_zp(wb2)
    wb2.close()
    if mz_agg:
        result['Медиана_ЗП_agg'] = mz_agg
        write_mediana_agg_sheet(mz_agg)   # → лист в database.xlsx для FastAPI

    # Вычисляем median_zp_computed (для карточки "Средняя зарплата (медиана)" на merger.kz)
    # Правило статусов: 2026 → активные, 2024/2025 → активные + "Завершен"
    mz_computed = compute_median_zp(DATABASE_XLSX)
    if mz_computed:
        result['median_zp_computed'] = mz_computed

    wb.close()

    # ── Монолитный database.json (оставляем как fallback для совместимости) ──
    with open(DATABASE_JSON, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, default=str)
    size_mb = os.path.getsize(DATABASE_JSON) / 1024 / 1024
    print(f'  ✓ database.json сохранён ({size_mb:.1f} МБ, {len(result)} листов)')

    # ── Разбиваем на 2 файла для быстрой загрузки на merger.kz ──
    split_into_main_and_msp(result)

    return True


def split_into_main_and_msp(result: dict):
    """Разбивает большой result на два файла:
       - db_main.json:  всё кроме 'msp' (быстро грузится при старте)
       - db_msp.json:   только 'msp'    (лениво, с прогресс-баром)

    Оставляем 'dzo' и прочие служебные секции в обоих файлах — их размер
    копеечный, а это даёт устойчивость если один из файлов не загрузился.
    """
    main_result = {k: v for k, v in result.items() if k not in MSP_ONLY_SECTIONS}
    msp_result  = {k: v for k, v in result.items() if k in MSP_ONLY_SECTIONS}

    # dzo-метаданные дублируем в msp-файл — они крошечные, но нужны для фильтрации
    for meta_key in ('dzo',):
        if meta_key in result and meta_key not in msp_result:
            msp_result[meta_key] = result[meta_key]

    with open(DB_MAIN_JSON, 'w', encoding='utf-8') as f:
        json.dump(main_result, f, ensure_ascii=False, default=str)
    main_mb = os.path.getsize(DB_MAIN_JSON) / 1024 / 1024
    print(f'  ✓ db_main.json сохранён ({main_mb:.1f} МБ, {len(main_result)} секций)')

    with open(DB_MSP_JSON, 'w', encoding='utf-8') as f:
        json.dump(msp_result, f, ensure_ascii=False, default=str)
    msp_mb = os.path.getsize(DB_MSP_JSON) / 1024 / 1024
    print(f'  ✓ db_msp.json сохранён ({msp_mb:.1f} МБ, {len(msp_result)} секций)')


# ─── Шаг 4: Git push на GitHub ─────────────────────────────────────────────

def git_push():
    print('\n[4] Отправка на GitHub...')

    remote_url = f'https://DarynZhamil:{GITHUB_TOKEN}@github.com/DarynZhamil/VideoforPBI.git'

    def run(cmd, **kwargs):
        return subprocess.run(cmd, cwd=REPO_DIR, capture_output=True, text=True, **kwargs)

    # Настройка remote с токеном
    run(['git', 'remote', 'set-url', 'origin', remote_url])

    # Добавляем файлы (включая разбитые db_main.json + db_msp.json)
    run(['git', 'add', 'database.xlsx', 'database.json',
         'db_main.json', 'db_msp.json'])

    # Проверяем есть ли изменения
    status = run(['git', 'status', '--porcelain'])
    if not status.stdout.strip():
        print('  Файлы не изменились — push не нужен')
        return True

    # Коммит
    now = datetime.now().strftime('%d.%m.%Y %H:%M')
    msg = f'Обновление базы данных {now}'
    result = run(['git', 'commit', '-m', msg,
                  '--author', 'DarynZhamil <daryn@kmg.kz>'])
    if result.returncode != 0:
        print(f'  Ошибка коммита: {result.stderr}')
        return False

    # Push (force — данные всегда берём локальные)
    result = run(['git', 'push', '--force', 'origin', 'main'])
    if result.returncode != 0:
        print(f'  Ошибка push: {result.stderr}')
        return False

    print(f'  ✓ Успешно загружено на GitHub!')
    print(f'  Коммит: {msg}')
    return True


# ─── Главная функция ────────────────────────────────────────────────────────

def main():
    print('=' * 60)
    print('  ОБНОВЛЕНИЕ БАЗЫ ДАННЫХ КМГ НДО')
    print(f'  {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
    print('=' * 60)

    # Шаг 1: Мониторинг из Word
    print('\n[1] Чтение Word-файлов мониторинга...')
    monitoring_df = build_monitoring_df()

    # Шаг 2: Обновление Excel
    if not update_monitoring_sheet(monitoring_df):
        print('\n  Ошибка на шаге 2 — прерываю')
        _pause()
        return

    # Шаг 3: Конвертация в JSON
    if not xlsx_to_json():
        print('\n  Ошибка на шаге 3 — прерываю')
        _pause()
        return

    # Шаг 4: GitHub
    git_push()

    print('\n' + '=' * 60)
    print('  ГОТОВО! Приложение получит новые данные')
    print('  в течение 24 часов (или сразу после перезагрузки)')
    print('=' * 60)
    _pause()


def _pause():
    try:
        input('\nНажмите Enter для выхода...')
    except (EOFError, KeyboardInterrupt):
        pass


if __name__ == '__main__':
    main()
